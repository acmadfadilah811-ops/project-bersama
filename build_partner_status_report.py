from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path(r"C:\bintang-project\deliverables\Laporan_Kondisi_Aplikasi_Mitra_Advertising_Tim_IT_Brandy.docx")

NAVY = "17365D"
BLUE = "1F4E79"
TEAL = "0F6B78"
ORANGE = "E67E22"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E9F5F3"
LIGHT_ORANGE = "FDF2E9"
GRAY = "5B6573"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 567)))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, size=None, color=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Aptos"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Aptos")
    return run


def add_text(doc, text, style=None, space_after=6, color=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    add_run(p, text, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    add_run(p, text)
    return p


def add_section_title(doc, number, title, intro=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, number + "  ", 11, ORANGE, bold=True)
    add_run(p, title, 16, NAVY, bold=True)
    if intro:
        add_text(doc, intro, space_after=7, color=GRAY)


def add_status_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, value in zip(header.cells, headers):
        shade(cell, NAVY)
        set_cell_margins(cell, 120, 130, 120, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, value, 9, WHITE, bold=True)
        set_cell_border(cell, top={"val":"single","sz":"4","color":NAVY}, bottom={"val":"single","sz":"4","color":NAVY})
    for idx, row_values in enumerate(rows):
        row = table.add_row()
        for c_idx, (cell, value) in enumerate(zip(row.cells, row_values)):
            shade(cell, WHITE if idx % 2 == 0 else LIGHT_GRAY)
            set_cell_margins(cell, 125, 140, 125, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if c_idx == 0:
                add_run(p, value, 9, NAVY, bold=True)
            elif c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, value, 9, TEAL, bold=True)
            else:
                add_run(p, value, 9, GRAY)
            set_cell_border(cell, bottom={"val":"single","sz":"3","color":"D9E2F3"})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [16.5])
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 160, 190, 160, 190)
    set_cell_border(cell,
                    top={"val":"single","sz":"8","color":ORANGE},
                    left={"val":"single","sz":"8","color":ORANGE},
                    bottom={"val":"single","sz":"8","color":ORANGE},
                    right={"val":"single","sz":"8","color":ORANGE})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, 10, NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    add_run(p2, text, 9, GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    add_run(p, "Tim IT Brandy | Laporan kondisi aplikasi untuk mitra advertising | 01 Agustus 2026", 8, GRAY)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("263238")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [("Title", 29, NAVY), ("Subtitle", 13, GRAY), ("Heading 1", 16, NAVY), ("Heading 2", 12, BLUE)]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True if "Heading" in style_name or style_name == "Title" else False


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    core = doc.core_properties
    core.title = "Laporan Kondisi Aplikasi - Mitra Advertising"
    core.subject = "Ringkasan kondisi aplikasi dan kesiapan operasional"
    core.author = "Tim IT Brandy"
    core.company = "Brandy"
    core.comments = "Dokumen untuk koordinasi mitra advertising"

    # Cover page - customer_pack inspired header.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    add_run(p, "TIM IT BRANDY", 11, ORANGE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Laporan Kondisi", 30, NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "Aplikasi & Kesiapan Operasional", 30, NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "Untuk Mitra Advertising", 15, TEAL, bold=True)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    set_table_widths(band, [16.5])
    cell = band.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    set_cell_margins(cell, 175, 210, 175, 210)
    set_cell_border(cell, left={"val":"single","sz":"16","color":ORANGE})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    add_run(p, "Ringkasan tujuan", 10, NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, "Menyampaikan kondisi aplikasi berdasarkan pemeriksaan operasional terkini, area yang telah siap digunakan, serta fokus validasi sebelum penerapan skala penuh.", 10, GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    set_table_widths(meta, [4.2, 12.3])
    metadata = [("Disusun oleh", "Tim IT Brandy"), ("Tanggal laporan", "01 Agustus 2026"), ("Klasifikasi", "Dokumen koordinasi operasional")]
    for row, (label, value) in zip(meta.rows, metadata):
        for cell in row.cells:
            set_cell_margins(cell, 110, 130, 110, 130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(row.cells[0], NAVY)
        shade(row.cells[1], LIGHT_GRAY)
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, label, 9, WHITE, bold=True)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, value, 9, GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Catatan ruang lingkup", 9, NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    add_run(p, "Laporan ini merupakan ringkasan teknis-operasional. Laporan tidak menggantikan audit independen, pengujian beban formal, atau persetujuan bisnis dan keuangan dari masing-masing pihak.", 8.5, GRAY, italic=True)

    doc.add_page_break()

    # Page 2.
    add_section_title(doc, "01", "Ringkasan Eksekutif")
    add_text(doc, "Aplikasi telah memiliki fondasi operasional yang aktif pada lingkungan VPS, dengan layanan aplikasi, basis data PostgreSQL, penyimpanan media, serta modul transaksi utama yang menggunakan data aplikasi. Berdasarkan pemeriksaan saat ini, aplikasi direkomendasikan untuk tahap UAT (User Acceptance Test) terkontrol bersama pengguna bisnis.", space_after=7)
    add_callout(doc, "Status rekomendasi saat ini: siap untuk UAT terkontrol", "Tahap ini ditujukan untuk mengonfirmasi alur bisnis nyata, kesiapan pengguna, dan kontrol operasional sebelum penerapan dengan domain permanen serta penggunaan skala penuh.", LIGHT_TEAL)

    add_section_title(doc, "02", "Kondisi Layanan dan Data", "Ringkasan berikut memisahkan komponen yang telah tersedia dari fokus pengujian lanjutan.")
    add_status_table(doc,
        ["Area", "Kondisi", "Catatan operasional"],
        [
            ("Layanan aplikasi", "Aktif", "Layanan web dan aplikasi berjalan pada lingkungan VPS dan dapat mendukung akses operasional."),
            ("Basis data", "Terhubung", "Aplikasi menggunakan PostgreSQL sebagai basis data utama untuk data transaksi dan modul pendukung."),
            ("Kasir & shift", "Data-backed", "Transaksi POS, ringkasan shift, dan data penjualan menggunakan penyimpanan aplikasi, bukan data tampilan statis."),
            ("Laporan", "Data-backed", "Laporan penjualan, pembelian, dan pembayaran telah diarahkan untuk menggunakan data operasional yang tersedia."),
            ("Penyimpanan gambar", "Aktif", "Media produk disimpan pada penyimpanan lokal VPS dan disajikan melalui layanan aplikasi."),
            ("Kapasitas VPS", "Memadai untuk UAT", "Konfigurasi saat ini 8 vCPU, memori sekitar 16 GB, dan ruang disk bebas sekitar 93 GB; tetap perlu pemantauan saat volume meningkat."),
        ],
        [3.5, 3.0, 10.0]
    )

    add_section_title(doc, "03", "Nilai Operasional untuk Mitra")
    add_bullet(doc, "Alur transaksi utama telah berada pada satu basis data operasional sehingga pencatatan dan pelaporan dapat ditelusuri dari sumber yang sama.")
    add_bullet(doc, "Penyimpanan media lokal memastikan gambar produk tetap dapat disajikan tanpa ketergantungan pada tautan penyimpanan eksternal selama operasional VPS berjalan normal.")
    add_bullet(doc, "Pemantauan tahap UAT dapat difokuskan pada pengalaman pengguna, ketepatan data, serta prosedur kerja mitra sebelum trafik dan penggunaan diperluas.")

    doc.add_page_break()

    # Page 3.
    add_section_title(doc, "04", "Bagian Akuntansi")
    add_text(doc, "Struktur akuntansi telah terhubung dengan PostgreSQL pada lingkungan aplikasi. Pemeriksaan data menunjukkan tabel akuntansi tersedia dan daftar akun telah terbentuk untuk mendukung konfigurasi serta pengembangan proses pencatatan keuangan.", space_after=6)
    add_status_table(doc,
        ["Komponen", "Kondisi", "Fokus UAT"],
        [
            ("Struktur data akuntansi", "Tersedia", "Konfirmasi akses, struktur akun, dan konsistensi data master bersama pengguna keuangan."),
            ("Daftar akun", "Terbentuk", "Validasi pemetaan akun sesuai kebijakan pencatatan dan kebutuhan pelaporan mitra."),
            ("Pencatatan jurnal", "Perlu validasi alur", "Uji transaksi contoh dari POS dan pembelian hingga terbentuk jurnal yang sesuai prosedur bisnis."),
            ("Rekonsiliasi", "Tahap UAT", "Bandingkan transaksi sumber, saldo kas, dan keluaran laporan agar selaras sebelum penggunaan skala penuh."),
        ],
        [4.0, 3.2, 9.3]
    )
    add_callout(doc, "Catatan profesional", "Kesiapan struktur dan konektivitas data adalah fondasi yang baik. Sebelum penetapan go-live penuh, Tim IT Brandy merekomendasikan pengujian bersama bagian keuangan untuk menyetujui pemetaan akun, bukti jurnal, serta rekonsiliasi transaksi operasional.", LIGHT_ORANGE)

    add_section_title(doc, "05", "Layanan WhatsApp")
    add_text(doc, "Konektivitas layanan WhatsApp melalui Evolution API telah tersedia pada lingkungan operasional dan instance integrasi berada dalam kondisi terhubung. Pengaturan WhatsApp pada antarmuka kasir saat ini sengaja dikunci atau disembunyikan dari alur pengguna harian.", space_after=6)
    add_text(doc, "Penguncian ini merupakan langkah tata kelola untuk memastikan konfigurasi, pengiriman resi, dan penggunaan nomor layanan divalidasi melalui proses yang terkontrol. Dengan pendekatan tersebut, akses konfigurasi tidak berubah tanpa otorisasi dan pengalaman pengguna kasir tetap terjaga.", space_after=7)
    add_status_table(doc,
        ["Aspek", "Kondisi", "Langkah berikutnya"],
        [
            ("Koneksi layanan", "Terhubung", "Lanjutkan uji pesan dan notifikasi menggunakan skenario pengguna yang disetujui."),
            ("Pengaturan pada kasir", "Terkunci", "Dibuka kembali setelah tata kelola konfigurasi dan alur administrasi disepakati."),
            ("Resi melalui WhatsApp", "UAT terkontrol", "Validasi nomor penerima, isi resi, status pengiriman, dan penanganan kegagalan."),
        ],
        [4.0, 3.2, 9.3]
    )

    doc.add_page_break()

    # Page 4.
    add_section_title(doc, "06", "Fokus Sebelum Penerapan Skala Penuh")
    add_text(doc, "Tahap berikut difokuskan pada penguatan operasional. Kegiatan ini bersifat preventif dan ditujukan agar aplikasi dapat dioperasikan secara konsisten saat penggunaan meningkat.", space_after=7)
    add_status_table(doc,
        ["Prioritas", "Kegiatan", "Hasil yang dituju"],
        [
            ("Tinggi", "Domain permanen dan Cloudflare Tunnel bernama", "Akses publik stabil, DNS dan HTTPS dikelola pada konfigurasi produksi yang terdokumentasi."),
            ("Tinggi", "Uji pemulihan cadangan PostgreSQL", "Bukti bahwa data dapat dipulihkan, disertai retensi dan salinan cadangan di lokasi terpisah."),
            ("Tinggi", "Pengamanan kredensial dan server", "Izin file konfigurasi, pengelolaan secret, serta akses administratif ditinjau sebelum go-live."),
            ("Menengah", "Pemantauan dan prosedur rollback", "Tim memiliki log, indikator layanan, kontak respons, dan langkah pemulihan yang jelas."),
            ("Menengah", "UAT end-to-end", "Pengguna bisnis menyetujui transaksi POS, shift, laporan, akuntansi, dan alur notifikasi yang relevan."),
            ("Menengah", "Uji kapasitas", "Kinerja dasar diukur pada skenario penggunaan aktual sebagai acuan peningkatan sumber daya VPS."),
        ],
        [2.2, 6.0, 8.3]
    )

    add_section_title(doc, "07", "Kesimpulan")
    add_text(doc, "Aplikasi berada pada posisi yang positif untuk memasuki UAT terkontrol. Komponen inti layanan, penyimpanan data, kasir, shift, laporan, dan media telah tersedia pada lingkungan operasional. Area akuntansi dan WhatsApp memiliki jalur penguatan yang jelas: akuntansi dilanjutkan melalui validasi jurnal dan rekonsiliasi, sedangkan konfigurasi WhatsApp tetap dikunci untuk menjaga tata kelola sampai pengujian pengiriman selesai.", space_after=7)
    add_callout(doc, "Rekomendasi Tim IT Brandy", "Laksanakan UAT bersama perwakilan operasional dan keuangan, dokumentasikan hasilnya, lalu selesaikan domain permanen, pemulihan cadangan, pengamanan server, dan prosedur operasional sebelum keputusan go-live penuh.", LIGHT_TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Disiapkan oleh", 9, GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Tim IT Brandy", 11, NAVY, bold=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
